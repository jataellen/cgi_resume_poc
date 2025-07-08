# from spire.doc import *
# from spire.doc.common import *
# from docx import Document as DocxDocument
import io
import os
import traceback
from src.logs_manager import log

# Import required libraries for PDF conversion
try:
    from docx import Document as DocxDocument
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
    from reportlab.lib.units import inch
    CONVERSION_AVAILABLE = True
except ImportError as e:
    log(f"PDF conversion libraries not available: {e}")
    CONVERSION_AVAILABLE = False


def convert_docx_to_pdf(uploaded_file, temp_file_path, pdf_file_path):
    """
    Convert a DOCX file to PDF using python-docx and reportlab
    
    Args:
        uploaded_file: File object with getbuffer() method
        temp_file_path: Path to save temporary DOCX file
        pdf_file_path: Path to save output PDF file
    
    Returns:
        str: Path to the created PDF file
    """
    if not CONVERSION_AVAILABLE:
        raise Exception("PDF conversion libraries not installed. Please install: pip install python-docx reportlab")
    
    try:
        # Save the uploaded DOCX file temporarily
        with open(temp_file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        log(f"DOCX file saved as: {temp_file_path}")

        # Load the DOCX document
        doc = DocxDocument(temp_file_path)
        
        # Create PDF document
        pdf_doc = SimpleDocTemplate(
            pdf_file_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            alignment=TA_LEFT
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=12,
            alignment=TA_LEFT
        )
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Check if it looks like a heading (all caps, short, etc.)
                if len(text) < 100 and (text.isupper() or any(run.bold for run in paragraph.runs)):
                    para = Paragraph(text, heading_style)
                else:
                    # Clean text for PDF
                    cleaned_text = text.replace('\u2022', '•').replace('\u2013', '-').replace('\u2014', '-')
                    para = Paragraph(cleaned_text, normal_style)
                story.append(para)
        
        # Process tables
        for table in doc.tables:
            # Add table content as paragraphs (simple approach)
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    para = Paragraph(row_text, normal_style)
                    story.append(para)
            story.append(Spacer(1, 0.2*inch))
        
        # Build PDF
        pdf_doc.build(story)
        log(f"Successfully converted DOCX to PDF: {pdf_file_path}")
        
        return pdf_file_path
        
    except Exception as e:
        log(f"Error in convert_docx_to_pdf: {str(e)}")
        log(traceback.format_exc())
        # Clean up any partial files
        for path in [temp_file_path, pdf_file_path]:
            if os.path.exists(path):
                try:
                    os.remove(path)
                    log(f"Cleaned up partial file: {path}")
                except Exception:
                    pass
        raise Exception(f"DOCX to PDF conversion failed: {str(e)}")


def change_all_fonts(uploaded_file):
    """
    Change all fonts in a DOCX document to Times New Roman
    
    Args:
        uploaded_file: File object with getbuffer() method
    
    Returns:
        ModifiedFile object with updated content
    """
    if not CONVERSION_AVAILABLE:
        log("Font changing not available - python-docx not installed")
        return uploaded_file
        
    try:
        # Create a BytesIO object from the uploaded file's buffer
        file_buffer = io.BytesIO(uploaded_file.getbuffer())

        # Load the document with python-docx
        doc = DocxDocument(file_buffer)

        # Change font for all runs in all paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Times New Roman"

        # Change font in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "Times New Roman"

        # Change font in headers and footers
        for section in doc.sections:
            for header in section.header.paragraphs:
                for run in header.runs:
                    run.font.name = "Times New Roman"

            for footer in section.footer.paragraphs:
                for run in footer.runs:
                    run.font.name = "Times New Roman"

        # Save the modified document to a new BytesIO object
        output_buffer = io.BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)  # Reset the buffer position to the beginning

        # Create a new FileStorage-like object with the modified document
        class ModifiedFile:
            def __init__(self, file_buffer, orig_file):
                self.buffer = file_buffer
                self.type = orig_file.type

            def getbuffer(self):
                return self.buffer.getvalue()

            def read(self):
                return self.buffer.read()

        log("Font changes applied successfully")
        return ModifiedFile(output_buffer, uploaded_file)
        
    except Exception as e:
        log(f"Error in change_all_fonts: {str(e)}")
        log(traceback.format_exc())
        # If font changing fails, return the original file
        log("Returning original file without font changes")
        return uploaded_file


def convert_to_pdf(uploaded_file, file_id):
    """
    Convert an uploaded file to PDF format. Handles both DOCX and PDF files.

    Args:
        uploaded_file: The uploaded file object
        file_id: A unique identifier for temp file naming

    Returns:
        str: Path to the PDF file

    Raises:
        Exception: If conversion fails
    """
    # Define file paths
    base_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    temp_file_path = os.path.join(base_path, f"temp_{file_id}.docx")
    pdf_file_path = os.path.join(base_path, f"temp_{file_id}.pdf")

    log(f"Preparing to process file with id {file_id}")

    try:
        if (
            uploaded_file.type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            log(f"Detected DOCX file: {uploaded_file.name}")

            try:
                # Change all fonts to Times New Roman
                modified_file = change_all_fonts(uploaded_file)
                log("Font changes applied")
            except Exception as e:
                log(
                    f"Warning: Font changing failed: {str(e)}. Continuing with original file."
                )
                modified_file = uploaded_file

            try:
                # Convert to PDF
                pdf_file_path = convert_docx_to_pdf(
                    modified_file, temp_file_path, pdf_file_path
                )
                log(f"Successfully converted to PDF: {pdf_file_path}")
            except Exception as e:
                log(f"Error in DOCX to PDF conversion: {str(e)}")
                raise Exception(f"Failed to convert DOCX to PDF: {str(e)}")
        else:
            # Assume it's already a PDF
            log(f"Detected PDF file: {uploaded_file.name}")
            with open(pdf_file_path, "wb") as f:
                f.write(uploaded_file.read())
            log(f"Saved PDF file: {pdf_file_path}")

        # Verify the file was created
        if not os.path.exists(pdf_file_path):
            raise Exception(f"PDF file was not created at {pdf_file_path}")

        # Clean up temporary DOCX file if it exists
        if os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
                log(f"Cleaned up temporary DOCX file: {temp_file_path}")
            except Exception:
                pass

        return pdf_file_path

    except Exception as e:
        log(f"Error in convert_to_pdf: {str(e)}")
        log(traceback.format_exc())

        # Clean up any temporary files
        for path in [temp_file_path, pdf_file_path]:
            if os.path.exists(path):
                try:
                    os.remove(path)
                    log(f"Cleaned up temporary file: {path}")
                except Exception:
                    pass

        raise Exception(f"Failed to process file: {str(e)}")