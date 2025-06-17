import sys
import os
# Add current directory to path for local imports
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, Header, Form
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List
import os
import shutil
import uuid
from datetime import datetime
import asyncio
from concurrent.futures import ThreadPoolExecutor
import traceback
import tempfile
import sys
from dotenv import load_dotenv
from supabase import create_client, Client
import jwt

# Load environment variables
load_dotenv()

# Set working directory to the script's directory for proper file paths
script_dir = os.path.dirname(os.path.abspath(__file__))
os.chdir(script_dir)
print(f"Working directory set to: {os.getcwd()}")

# Verify required files exist
required_files = [
    'data/experience_schema.json',
    'data/json_schema.json', 
    'data/all_schemas.json',
    'data/prompts.py'
]

missing_files = []
for file_path in required_files:
    if not os.path.exists(file_path):
        missing_files.append(file_path)

if missing_files:
    print("❌ Missing required files:")
    for file_path in missing_files:
        print(f"   - {file_path}")
    print("\\nPlease make sure you're running the app from the correct directory")
    print("and that all data files are present.")
    sys.exit(1)

# Initialize FastAPI app
app = FastAPI(title="ResumeGenie API", version="1.0.0")

# Get environment variables
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY")
FRONTEND_URL = os.getenv("FRONTEND_URL", "http://localhost:3000")
PORT = int(os.getenv("PORT", "8000"))

# Initialize Supabase client
supabase: Client = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY) if SUPABASE_URL and SUPABASE_SERVICE_KEY else None

# Configure CORS for React frontend
allowed_origins = [
    FRONTEND_URL,
    "http://localhost:3000",
    "https://cgi-resumegenie-staging.netlify.app",
    "https://cgi-resumegenie.netlify.app"  # Production URL
]

# Filter out None/empty values
allowed_origins = [origin for origin in allowed_origins if origin]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Thread pool for CPU-bound tasks
executor = ThreadPoolExecutor(max_workers=4)

# Store upload sessions
upload_sessions = {}

# Auth dependency
async def get_current_user(authorization: Optional[str] = Header(None)):
    """Verify JWT token from Supabase"""
    
    # Require Supabase configuration
    if not SUPABASE_URL or not supabase:
        raise HTTPException(status_code=500, detail="Authentication service not configured")
    
    # Require authorization header
    if not authorization:
        raise HTTPException(status_code=401, detail="Authorization header missing")
    
    # Extract token from authorization header
    if not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Invalid authorization header format")
    
    token = authorization.split(" ")[1]
    
    try:
        # Use Supabase client to verify the token and get user
        from supabase import Client
        
        # Create a new client with the user's token
        auth_client = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)
        
        # Get user from token
        user_response = auth_client.auth.get_user(token)
        
        if not user_response or not user_response.user:
            raise HTTPException(status_code=401, detail="Invalid authentication token")
            
        user = user_response.user
        return {
            "id": user.id, 
            "email": user.email, 
            "mode": "supabase", 
            "user_data": user
        }
        
    except Exception as e:
        print(f"Supabase auth failed: {e}")
        raise HTTPException(status_code=401, detail="Authentication failed")

# Pydantic models for request/response
class UploadResponse(BaseModel):
    session_id: str
    message: str
    filename: str

class ProcessStatus(BaseModel):
    status: str  # "processing", "completed", "error"
    logs: List[str]
    progress: int  # 0-100
    download_url: Optional[str] = None
    error: Optional[str] = None

# Create necessary directories
os.makedirs("uploads", exist_ok=True)
os.makedirs("outputs", exist_ok=True)

# Progress tracking helper
def update_session_progress(session_id, progress):
    """Update progress for a session"""
    if session_id in upload_sessions:
        # Convert 0-1 range to 0-100 if needed
        if progress <= 1:
            progress_percent = int(progress * 100)
        else:
            progress_percent = int(progress)
        upload_sessions[session_id]["progress"] = min(progress_percent, 100)

def simple_process_resume(file_path, session_id, format_type="Developer"):
    """Simplified resume processing without streamlit dependencies"""
    try:
        update_session_progress(session_id, 20)
        
        # For now, this is a placeholder that copies the file
        # In a real implementation, you'd add your resume processing logic here
        update_session_progress(session_id, 50)
        
        # Simulate processing by copying the file with a new name
        output_file = "updated_resume.docx"
        
        if file_path.endswith('.docx'):
            shutil.copy2(file_path, output_file)
        else:
            # Convert PDF to DOCX would go here
            # For now, create a simple placeholder
            from docx import Document
            doc = Document()
            doc.add_paragraph(f"Resume processing completed for {format_type} format.")
            doc.add_paragraph("Note: Full processing temporarily disabled.")
            doc.save(output_file)
        
        update_session_progress(session_id, 100)
        
        if session_id in upload_sessions:
            upload_sessions[session_id]["logs"].append("Resume processing completed successfully!")
        
        return output_file
        
    except Exception as e:
        if session_id in upload_sessions:
            upload_sessions[session_id]["logs"].append(f"Processing error: {str(e)}")
        raise e

@app.get("/")
async def root():
    return {
        "message": "ResumeGenie API is running", 
        "version": "1.0.0",
        "status": "healthy",
        "working_directory": os.getcwd()
    }

@app.post("/api/upload", response_model=UploadResponse)
async def upload_resume(file: UploadFile = File(...), current_user=Depends(get_current_user)):
    """Upload a resume file (PDF or DOCX) for processing"""
    
    # Validate file type
    if not file.filename.endswith(('.pdf', '.docx', '.doc')):
        raise HTTPException(status_code=400, detail="Only PDF, DOCX, and DOC files are allowed")
    
    # Generate unique session ID
    session_id = str(uuid.uuid4())
    
    # Save uploaded file
    upload_path = f"uploads/{session_id}_{file.filename}"
    try:
        with open(upload_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")
    
    # Initialize session
    upload_sessions[session_id] = {
        "status": "uploaded",
        "filename": file.filename,
        "upload_path": upload_path,
        "output_path": None,
        "logs": [f"File received: {file.filename}"],
        "progress": 10,
        "created_at": datetime.now(),
        "error": None,
        "user_id": current_user.get("id", "unknown")
    }
    
    # Start processing in background
    asyncio.create_task(process_resume_async(session_id))
    
    return UploadResponse(
        session_id=session_id,
        message="File uploaded successfully. Processing will begin shortly.",
        filename=file.filename
    )

async def process_resume_async(session_id: str):
    """Process resume asynchronously"""
    session = upload_sessions[session_id]
    
    try:
        # Update status
        session["status"] = "processing"
        session["logs"].append("Initializing document processing...")
        session["progress"] = 20
        
        # Get file info
        upload_path = session["upload_path"]
        original_filename = session["filename"]
        
        session["logs"].append("Processing resume...")
        session["progress"] = 30
        
        # Run the simplified resume processing in thread pool
        loop = asyncio.get_event_loop()
        
        await loop.run_in_executor(
            executor,
            simple_process_resume,
            upload_path,
            session_id,
            "Developer"  # Default format
        )
        
        # Handle output file
        output_filename = os.path.splitext(original_filename)[0] + "_updated.docx"
        temp_output = "updated_resume.docx"
        
        if os.path.exists(temp_output):
            final_output_path = f"outputs/{session_id}_{output_filename}"
            shutil.move(temp_output, final_output_path)
            session["output_path"] = final_output_path
            session["logs"].append("Resume processing completed successfully!")
            session["progress"] = 100
            session["status"] = "completed"
        else:
            raise Exception("Output file was not generated")
            
    except Exception as e:
        session["status"] = "error"
        error_msg = f"Processing error: {str(e)}"
        session["logs"].append(error_msg)
        session["error"] = str(e)
        session["progress"] = 0
        
        # Log full traceback for debugging
        print(f"Error processing session {session_id}: {str(e)}")

@app.get("/api/status/{session_id}", response_model=ProcessStatus)
async def get_process_status(session_id: str, current_user=Depends(get_current_user)):
    """Get the current status of resume processing"""
    
    if session_id not in upload_sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    session = upload_sessions[session_id]
    
    # Verify user owns this session
    user_id = current_user.get("id")
    if session.get("user_id") != user_id:
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Prepare download URL if completed
    download_url = None
    if session["status"] == "completed" and session["output_path"]:
        download_url = f"/api/download/{session_id}"
    
    return ProcessStatus(
        status=session["status"],
        logs=session["logs"],
        progress=session["progress"],
        download_url=download_url,
        error=session.get("error")
    )

@app.get("/api/download/{session_id}")
async def download_resume(session_id: str, current_user=Depends(get_current_user)):
    """Download the processed resume"""
    
    if session_id not in upload_sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    session = upload_sessions[session_id]
    
    # Verify user owns this session
    user_id = current_user.get("id")
    if session.get("user_id") != user_id:
        raise HTTPException(status_code=403, detail="Access denied")
    
    if session["status"] != "completed" or not session["output_path"]:
        raise HTTPException(status_code=400, detail="Resume processing not completed")
    
    output_path = session["output_path"]
    if not os.path.exists(output_path):
        raise HTTPException(status_code=404, detail="Output file not found")
    
    # Extract original filename for download
    original_filename = session["filename"]
    download_filename = os.path.splitext(original_filename)[0] + "_updated.docx"
    
    return FileResponse(
        path=output_path,
        filename=download_filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.delete("/api/session/{session_id}")
async def cleanup_session(session_id: str, current_user=Depends(get_current_user)):
    """Clean up session data and files"""
    
    if session_id not in upload_sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    session = upload_sessions[session_id]
    
    # Verify user owns this session
    user_id = current_user.get("id")
    if session.get("user_id") != user_id:
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Clean up files
    if session.get("upload_path") and os.path.exists(session["upload_path"]):
        os.remove(session["upload_path"])
    
    if session.get("output_path") and os.path.exists(session["output_path"]):
        os.remove(session["output_path"])
    
    # Remove session
    del upload_sessions[session_id]
    
    return {"message": "Session cleaned up successfully"}

@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "active_sessions": len(upload_sessions),
        "working_directory": os.getcwd(),
        "streamlit_removed": True
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=PORT, reload=True)