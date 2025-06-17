from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def replace_text_in_docx(doc: Document, replacements: list):
    for key, value in replacements:
        is_missing = (
            value is None or 
            (isinstance(value, str) and not value.strip()) or 
            (isinstance(value, list) and not any(item.strip() for item in value))
        )

        # If non-list and missing, just replace it with a message
        if is_missing and not isinstance(value, list):
            value = f"[Missing field: {key.strip('{}')}]"

        def process_paragraphs(paragraphs):
            for paragraph in paragraphs:
                if key in paragraph.text:
                    if isinstance(value, list):
                        # Remove the paragraph with the placeholder
                        parent = paragraph._element.getparent()
                        idx = parent.index(paragraph._element)
                        parent.remove(paragraph._element)

                        if is_missing:
                            # Insert "Missing field" paragraph
                            para = doc.add_paragraph()
                            run = para.add_run(f"[Missing field: {key.strip('{}')}]")
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 0, 0)
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            parent.insert(idx, para._element)
                        else:
                            # Insert bullet points
                            for item in value:
                                if item.strip():
                                    bullet_para = doc.add_paragraph(item.strip(), style="ListBullet")
                                    parent.insert(idx, bullet_para._element)
                                    idx += 1
                    else:
                        # Simple text replacement
                        paragraph.text = paragraph.text.replace(key, str(value))
                    return True
            return False

        # Try replacing in normal paragraphs
        if process_paragraphs(doc.paragraphs):
            continue

        # Try replacing inside table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if process_paragraphs(cell.paragraphs):
                        break
