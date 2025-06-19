from openai import AzureOpenAI
from langchain_openai import AzureChatOpenAI
from langchain.schema import SystemMessage, HumanMessage
from langchain_community.document_loaders import PyPDFLoader
from src.experience_chain import experience_chain
from src.logs_manager import log
from src.resume_generator import generate_resume
import os
import json
import datetime
from dotenv import load_dotenv
from utils.document_utils import *
from data.prompts import *
from src.cgi_experience_generator import generate_default_cgi_prompt

# Import additional RAG-related libraries
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import CharacterTextSplitter
from langchain_openai import AzureOpenAIEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
import traceback
from docx import Document
import re

# Import missing functions needed for resume_stream
from src.cgi_experience_generator import generate_cgi_experience

def generate_rag_job_description(llm, file_path, role_title):
    """
    Uses RAG to analyze an RFP document and generate a comprehensive job description.
    """
    try:
        log(f"Generating job description from RFP file using RAG for role: {role_title}")

        # Load the PDF
        loader = PyPDFLoader(file_path)
        pages = loader.load()
        pdf_text = "\n".join([doc.page_content for doc in pages])

        # Split text into chunks
        text_splitter = CharacterTextSplitter(chunk_size=3000, chunk_overlap=100)
        documents = text_splitter.split_text(pdf_text)

        # Ensure documents are clean and valid
        documents = [
            doc for doc in documents if doc and isinstance(doc, str) and doc.strip()
        ]

        if not documents:
            log("Warning: No valid text extracted from the document")
            return ""

        # Create embeddings with proper deployment
        embeddings = AzureOpenAIEmbeddings(
            model="text-embedding-3-large",
            azure_endpoint=os.environ["AZURE_OPENAI_EMBEDDINGS_ENDPOINT"],
            api_key=os.environ["AZURE_OPENAI_API_KEY"],
        )

        # Log for debugging
        log(f"Creating vector store with {len(documents)} document chunks")

        # Create vector store
        vector_store = FAISS.from_texts(documents, embeddings)
        retriever = vector_store.as_retriever()

        # Create prompt
        prompt = ChatPromptTemplate.from_messages([
            ("system", RFP_SP),
            ("human", RFP_HP),
        ])

        # Create and execute retrieval chain
        question_answer_chain = create_stuff_documents_chain(llm, prompt)
        chain = create_retrieval_chain(retriever, question_answer_chain)

        # Invoke with the specific role
        log(f"Executing RAG retrieval chain for {role_title}")
        response = chain.invoke({"input": role_title})
        job_description = response["answer"]

        log(f"Successfully generated job description from RFP ({len(job_description)} chars)")
        return job_description

    except Exception as e:
        log(f"Error generating job description from RFP: {str(e)}")
        log(traceback.format_exc())
        return ""
    
def extract_resume_data_from_docx(file_path):
    """
    Extract resume data from a DOCX file and return structured JSON
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        dict: Structured resume data
    """
    try:
        doc = Document(file_path)
        
        # Extract all text from the document
        full_text = []
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            full_text.append(paragraph.text.strip())
        
        resume_text = '\n'.join(full_text)
        
        # Basic structure
        structured_data = {
            "contact": {
                "name": "",
                "email": "",
                "phone": "",
                "location": ""
            },
            "profile": "",
            "skills": [],
            "experience": [],
            "education": [],
            "projects": [],
            "volunteer_experience": [],
            "certifications": []
        }
        
        # Enhanced text-based parsing
        lines = [line.strip() for line in resume_text.split('\n') if line.strip()]
        current_section = None
        current_entry = {}
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            
            # Detect contact information (usually at the top)
            if '@' in line and not structured_data["contact"]["email"]:
                structured_data["contact"]["email"] = line
            elif any(char in line for char in ['(', ')', '-']) and len([c for c in line if c.isdigit()]) >= 7:
                if not structured_data["contact"]["phone"]:
                    structured_data["contact"]["phone"] = line
            
            # Detect sections
            if any(keyword in line_lower for keyword in ['summary', 'profile', 'objective']):
                current_section = 'profile'
                continue
            elif any(keyword in line_lower for keyword in ['skills', 'technical skills', 'competencies']):
                current_section = 'skills'
                continue
            elif any(keyword in line_lower for keyword in ['experience', 'work experience', 'employment']):
                current_section = 'experience'
                continue
            elif any(keyword in line_lower for keyword in ['education', 'academic']):
                current_section = 'education'
                continue
            elif any(keyword in line_lower for keyword in ['projects', 'project experience']):
                current_section = 'projects'
                continue
            elif any(keyword in line_lower for keyword in ['volunteer', 'community']):
                current_section = 'volunteer_experience'
                continue
            elif any(keyword in line_lower for keyword in ['certification', 'certificates']):
                current_section = 'certifications'
                continue
            
            # Process content based on current section
            if current_section == 'profile':
                if structured_data["profile"]:
                    structured_data["profile"] += " " + line
                else:
                    structured_data["profile"] = line
            
            elif current_section == 'skills':
                # Try to parse skills (common formats: comma-separated, bullet points)
                if ',' in line:
                    skills = [skill.strip() for skill in line.split(',')]
                    structured_data["skills"].extend(skills)
                elif line.startswith('•') or line.startswith('-'):
                    structured_data["skills"].append(line.lstrip('•-').strip())
                else:
                    structured_data["skills"].append(line)
            
            elif current_section in ['experience', 'projects', 'volunteer_experience']:
                # Simple approach: treat each line as either a new entry or continuation
                if any(keyword in line_lower for keyword in ['company', 'position', 'title', 'role']):
                    if current_entry:
                        structured_data[current_section].append(current_entry)
                    current_entry = {"description": line, "details": []}
                else:
                    if current_entry:
                        current_entry["details"].append(line)
                    else:
                        current_entry = {"description": line, "details": []}
            
            elif current_section == 'education':
                structured_data["education"].append(line)
            
            elif current_section == 'certifications':
                structured_data["certifications"].append(line)
        
        # Add any remaining entry
        if current_entry and current_section in ['experience', 'projects', 'volunteer_experience']:
            structured_data[current_section].append(current_entry)
        
        # Try to extract name from the first few lines if not found
        if not structured_data["contact"]["name"]:
            for line in lines[:5]:
                if not any(char in line for char in ['@', '(', ')', '-']) and len(line.split()) <= 4:
                    structured_data["contact"]["name"] = line
                    break
        
        return structured_data
        
    except Exception as e:
        log(f"Error extracting resume data from DOCX: {str(e)}")
        return None

def extract_resume_data_with_llm(file_path, llm):
    """
    Extract resume data from DOCX using LLM for better parsing
    
    Args:
        file_path (str): Path to the DOCX file
        llm: Language model instance
        
    Returns:
        dict: Structured resume data
    """
    try:
        # First extract raw text
        doc = Document(file_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            full_text.append(paragraph.text.strip())
        
        resume_text = '\n'.join(full_text)
        
        # Use LLM to structure the data
        system_prompt = """
        You are an expert resume parser. Extract structured information from the given resume text.
        Return the data in the following JSON format:
        {
            "contact": {
                "name": "Full Name",
                "email": "email@example.com",
                "phone": "phone number",
                "location": "city, state"
            },
            "profile": "Professional summary or objective",
            "skills": ["skill1", "skill2", "skill3"],
            "experience": [
                {
                    "title": "Job Title",
                    "company": "Company Name",
                    "duration": "Start Date - End Date",
                    "responsibilities": ["responsibility1", "responsibility2"]
                }
            ],
            "education": [
                {
                    "degree": "Degree Type",
                    "field": "Field of Study",
                    "institution": "Institution Name",
                    "year": "Graduation Year"
                }
            ],
            "certifications": ["certification1", "certification2"]
        }
        
        Extract all available information. If a field is not present, use an empty string or empty array as appropriate.
        """
        
        human_prompt = f"Parse the following resume text and return structured JSON data:\n\n{resume_text}"
        
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=human_prompt)
        ]
        
        response = llm.invoke(messages)
        
        try:
            # Try to parse JSON from response
            json_match = re.search(r'```json\n(.*?)\n```', response.content, re.DOTALL)
            if json_match:
                result = json.loads(json_match.group(1))
                return result
            else:
                # Try to parse the entire response as JSON
                result = json.loads(response.content.strip())
                return result
        except Exception as parse_error:
            return extract_resume_data_from_docx(file_path)
            
    except Exception as e:
        return extract_resume_data_from_docx(file_path)

def generate_llm_content(
    llm,
    system_prompt,
    human_prompt_template,
    format_args=None,
    functions=None,
    extract_function_call=False,
):
    try:
        format_args = format_args or {}
        current_date = datetime.datetime.now().date()
        human_content = (
            human_prompt_template.format(**format_args)
            + f"\n\nCurrent Date: {current_date}"
        )

        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=human_content),
        ]

        response = llm.invoke(messages, functions=functions if functions else None)

        if extract_function_call and "function_call" in response.additional_kwargs:
            function_args = response.additional_kwargs["function_call"]["arguments"]
            return json.loads(function_args)
        else:
            return response.content.strip()
    except Exception as e:
        log(f"Error generating content: {str(e)}")
        return f"Error generating content: {str(e)}"

def generate_cgi_experience(llm, format_type, custom_role_title=None):
    """
    Uses the LLM to generate a default CGI experience entry.

    Args:
        llm: The LLM instance to use
        format_type: The selected resume format
        custom_role_title: Optional custom role title

    Returns:
        dict: A dictionary containing the CGI experience entry
    """
    try:
        log(f"Generating default CGI experience for {format_type} role")

        # Create the prompt for the LLM
        prompt = generate_default_cgi_prompt(format_type, custom_role_title)

        # Use simple system + human message approach
        messages = [
            SystemMessage(
                content="You are an expert resume writer who specializes in creating realistic and compelling work experiences for CGI consultants."
            ),
            HumanMessage(content=prompt),
        ]

        # Get response from LLM
        response = llm.invoke(messages)
        content = response.content.strip()

        # Extract JSON from the response
        try:
            # Try to parse the entire response as JSON
            cgi_experience = json.loads(content)
            log("Successfully generated CGI experience entry")
            return cgi_experience
        except json.JSONDecodeError:
            # If that fails, try to extract JSON from text
            log("Failed to parse response as JSON, attempting to extract JSON block")
            
            json_match = re.search(r"```json\n(.*?)\n```", content, re.DOTALL)
            if json_match:
                try:
                    cgi_experience = json.loads(json_match.group(1))
                    log("Successfully extracted JSON from markdown code block")
                    return cgi_experience
                except:
                    log("Failed to parse extracted JSON")

            # Last resort: Try to create a manual structure from the response
            log("Creating fallback CGI experience entry")

            # Get current date for start date
            current_date = datetime.datetime.now()
            start_date = f"{current_date.month:02d}/{current_date.year}"

            return {
                "cgi_client_or_sector": "Major Enterprise Client",
                "cgi_position_title": (
                    custom_role_title if custom_role_title else format_type
                ),
                "cgi_start_date": start_date,
                "cgi_end_date": "Present",
                "cgi_responsibilities": [
                    "Led development of enterprise solutions tailored to client business requirements.",
                    "Collaborated with cross-functional teams to implement system enhancements.",
                    "Provided technical expertise and guidance to ensure project success.",
                    "Conducted thorough analysis and testing to maintain quality standards.",
                ],
                "cgi_technologies": [
                    "Agile Methodology",
                    "Cloud Services",
                    "Enterprise Architecture",
                    "Business Intelligence",
                    "Data Management",
                ],
            }
    except Exception as e:
        log(f"Error generating CGI experience: {str(e)}")
        log(traceback.format_exc())

        # Fallback to a simple default
        current_date = datetime.datetime.now()
        start_date = f"{current_date.month:02d}/{current_date.year}"

        return {
            "cgi_client_or_sector": "Enterprise Client",
            "cgi_position_title": (
                custom_role_title if custom_role_title else format_type
            ),
            "cgi_start_date": start_date,
            "cgi_end_date": "Present",
            "cgi_responsibilities": [
                "Implemented solutions based on client requirements.",
                "Collaborated with cross-functional teams.",
                "Provided technical expertise throughout project lifecycle.",
                "Ensured high-quality deliverables.",
            ],
            "cgi_technologies": [
                "Agile",
                "Cloud Services",
                "Enterprise Software",
                "Data Management",
            ],
        }

def resume_stream(
    st,
    progress_bar,
    base_progress,
    file_progress_weight,
    file_path,
    selected_format,
    custom_role_title="",
    job_description="",
    rfp_file_path=None,
    include_default_cgi=False,  # New parameter
    custom_experiences=None,  # New parameter for custom experience entries
):
    load_dotenv()

    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    JSON_SCHEMA_PATH = os.path.join(BASE_DIR, "data", "json_schema.json")
    ALL_SCHEMAS_PATH = os.path.join(BASE_DIR, "data", "all_schemas.json")
    EXPERIENCE_SCHEMA_PATH = os.path.join(BASE_DIR, "data", "experience_schema.json")

    llm = AzureChatOpenAI(
        azure_endpoint=os.environ["AZURE_OPENAI_ENDPOINT"],
        api_key=os.environ["AZURE_OPENAI_API_KEY"],
        api_version="2024-12-01-preview",
        deployment_name="gpt-4o",
        model="gpt-4o",
    )

    # If RFP file is provided, generate a job description from it
    if rfp_file_path and os.path.exists(rfp_file_path):
        log(f"RFP file detected: {rfp_file_path}")
        # Determine role title to use for the RAG job description generation
        role_for_rag = (
            custom_role_title.strip() if custom_role_title else selected_format
        )

        # Generate job description from RFP
        rfp_job_description = generate_rag_job_description(
            llm, rfp_file_path, role_for_rag
        )

        # If successful and user didn't provide a job description, use the generated one
        if rfp_job_description and not (job_description and job_description.strip()):
            log("Using RFP-generated job description")
            job_description = rfp_job_description
        # If user provided a job description, combine it with the RFP-generated one
        elif rfp_job_description and job_description and job_description.strip():
            log("Combining user-provided job description with RFP-generated description")
            job_description = f"{job_description}\n\nAdditional requirements from RFP:\n{rfp_job_description}"

    # Update progress
    progress_bar.progress(base_progress + file_progress_weight * 0.1)

    loader = PyPDFLoader(file_path)
    pages = []
    # async for page in loader.alazy_load():
    #     pages.append(page)
    pages = loader.load()

    pdf_text = "\n".join([doc.page_content for doc in pages])

    current_date = datetime.datetime.now().date()

    try:
        with open(JSON_SCHEMA_PATH, "r") as file:
            json_schema = json.load(file)
    except FileNotFoundError:
        raise FileNotFoundError(f"JSON schema file not found at {JSON_SCHEMA_PATH}")
   
    # Generate structured data
    structured_data = generate_llm_content(
        llm=llm,
        system_prompt=STRUCTURED_DATA_SP,
        human_prompt_template=STRUCTURED_DATA_HP,
        format_args={
            "pdf_text": pdf_text,
            "json_input": json.dumps(json_schema, indent=2),
        },
        functions=[json_schema],
        extract_function_call=True,
    )

    log("Completed Structured Data")
    
    # If custom_experiences is provided, add them to the structured_data experience array
    if custom_experiences and len(custom_experiences) > 0:
        log(f"Adding {len(custom_experiences)} custom experience entries to structured data")
        
        # Ensure experience array exists
        if "experience" not in structured_data:
            structured_data["experience"] = []
        
        # Add custom experiences to the experience array
        for exp in custom_experiences:
            experience_entry = {
                "company": exp.get("company", ""),
                "position_title": exp.get("position_title", ""),
                "start_date": exp.get("start_date", ""),
                "end_date": exp.get("end_date", ""),
                "responsibilities": [exp.get("description", "")],
                "technologies": []  # Could be extracted from description if needed
            }
            structured_data["experience"].append(experience_entry)
        
        log(f"Added {len(custom_experiences)} custom experiences to structured data")
    
    # Update progress - 20% complete for this file
    progress_bar.progress(base_progress + file_progress_weight * 0.2)

    custom_role_title_string = (
        f"\nAnd considering the target role type of {custom_role_title}, aligning with the general role type, while being more general:\n"
        if custom_role_title
        else ""
    )

    # Generate an appropriate role title if not provided by user
    if not custom_role_title or not custom_role_title.strip():
        role_title = generate_llm_content(
            llm=llm,
            system_prompt=ROLE_TITLE_GEN_SP,
            human_prompt_template=ROLE_TITLE_GEN_HP,
            format_args={
                "structured_data": json.dumps(structured_data, indent=2),
                "custom_role_title_string": custom_role_title_string,
            },
        )
        log(f"Generated role title: {role_title}")
    else:
        role_title = custom_role_title.strip()
        log(f"Using provided role title: {role_title}")

    # Generate tailored profile based on job description if provided
    if job_description and job_description.strip():
        profile = generate_llm_content(
            llm=llm,
            system_prompt=TAILORED_SUMMARY_SP,
            human_prompt_template=TAILORED_SUMMARY_HP,
            format_args={
                "structured_data": structured_data,
                "job_description": job_description,
                "role": role_title,
            },
            functions=[json_schema],
        )
        log(f"Generated tailored profile with job description for {role_title}")
    else:
        profile = generate_llm_content(
            llm=llm,
            system_prompt=SUMMARY_SP,
            human_prompt_template=SUMMARY_HP,
            format_args={"structured_data": structured_data},
            functions=[json_schema],
        )
        log("Generated standard profile")

    # Generate years of experience
    years_exp = generate_llm_content(
        llm=llm,
        system_prompt=PROFILE_SP,
        human_prompt_template=PROFILE_HP,
        format_args={"profile": profile},
    )

    # Define function to call LLM with appropriate schema
    def call_llm(all_schemas, section, text_input=pdf_text, job_desc=""):
        sp_var = f"{section.upper()}_SP"
        hp_var = f"{section.upper()}_HP"

        # Only use tailored prompts if job description is provided and not empty
        if (
            job_desc
            and job_desc.strip()
            and f"TAILORED_{section.upper()}_SP" in globals()
        ):
            sp_var = f"TAILORED_{section.upper()}_SP"
            hp_var = f"TAILORED_{section.upper()}_HP"

        sp = eval(sp_var)
        hp = eval(hp_var)

        # Add job description to format args if available and needed
        format_args = {
            "text_input": text_input,
            "json_dump": json.dumps(all_schemas[section]["json_schema"], indent=2),
        }

        if (
            job_desc
            and job_desc.strip()
            and f"TAILORED_{section.upper()}_HP" in globals()
        ):
            format_args["job_description"] = job_desc
            format_args["role"] = role_title

        messages = [
            SystemMessage(content=sp),
            HumanMessage(content=hp.format(**format_args)),
        ]

        response = llm.invoke(messages, functions=[all_schemas[section]["json_schema"]])
        structured_data = response.additional_kwargs["function_call"]["arguments"]
        json_structured_data = json.loads(structured_data)

        return json_structured_data

    try:
        with open(ALL_SCHEMAS_PATH, "r") as file:
            all_schemas = json.load(file)
    except FileNotFoundError:
        raise FileNotFoundError(f"All schemas file not found at {ALL_SCHEMAS_PATH}")

    res_dict = dict()

    log("Loading...")
    # Process experience with job description if available and not empty
    if job_description and job_description.strip():
        res_dict["experience"] = call_tailored_experience_chain(
            pdf_text, job_description, role_title, llm
        )
        log(f"\t>> Completed key: experience (tailored for {role_title})")
    else:
        res_dict["experience"] = experience_chain(pdf_text, llm)
        log(f"\t>> Completed key: experience")

    # If include_default_cgi is True, generate a default CGI experience entry
    if include_default_cgi:
        log("Generating default CGI experience entry")
        default_cgi_exp = generate_cgi_experience(
            llm, selected_format, custom_role_title
        )

        # Add the default entry to the beginning of the CGI experience array
        if "cgi_experience" in res_dict["experience"]:
            # Check if there's already a non-empty cgi_experience
            if (
                res_dict["experience"]["cgi_experience"]
                and len(res_dict["experience"]["cgi_experience"]) > 0
            ):
                # Check if the first entry is a placeholder (client descriptor not provided)
                if res_dict["experience"]["cgi_experience"][0].get(
                    "cgi_client_or_sector"
                ) == "client descriptor not provided" or not res_dict["experience"][
                    "cgi_experience"
                ][
                    0
                ].get(
                    "cgi_responsibilities"
                ):
                    # Replace the placeholder with our generated experience
                    res_dict["experience"]["cgi_experience"][0] = default_cgi_exp
                    log("Replaced placeholder CGI experience with generated experience")
                else:
                    # Insert at the beginning if there's already valid content
                    res_dict["experience"]["cgi_experience"].insert(0, default_cgi_exp)
                    log("Added generated CGI experience at the beginning of existing experiences")
            else:
                # If the array is empty or None, initialize it with our experience
                res_dict["experience"]["cgi_experience"] = [default_cgi_exp]
                log("Created new CGI experience entry")

        log("Added default CGI experience entry")

    # Process other sections
    for key in all_schemas.keys():
        res_dict[key] = call_llm(all_schemas, key, pdf_text, job_description)
        log(f"\t>> Completed key: {key}")
    progress_bar.progress(base_progress + file_progress_weight * 0.4)

    # For testing
    data = {
        "structured_data": structured_data,
        "years_exp": years_exp,
        "profile": profile,
        "res_dict": res_dict,
        "job_description": job_description if job_description else "None provided",
    }
    with open("resume_data.json", "w") as json_file:
        json.dump(data, json_file, indent=4)

    print("Data saved to resume_data.json")

    # Pass job description, role title, and format type to generate_resume
    generate_resume(
        structured_data,
        years_exp,
        profile,
        res_dict,
        job_description,
        role_title,
        selected_format,
    )

def generate_llm_content(
    llm,
    system_prompt,
    human_prompt_template,
    format_args=None,
    functions=None,
    extract_function_call=False,
):
    try:
        format_args = format_args or {}
        current_date = datetime.datetime.now().date()
        human_content = (
            human_prompt_template.format(**format_args)
            + f"\n\nCurrent Date: {current_date}"
        )

        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=human_content),
        ]

        response = llm.invoke(messages, functions=functions if functions else None)

        if extract_function_call and "function_call" in response.additional_kwargs:
            function_args = response.additional_kwargs["function_call"]["arguments"]
            return json.loads(function_args)
        else:
            return response.content.strip()
    except Exception as e:
        log(f"Error generating content: {str(e)}")
        return f"Error generating content: {str(e)}"

def call_tailored_experience_chain(pdf_text, job_description, role, llm):
    """Modified experience chain that incorporates job description"""

    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    EXPERIENCE_SCHEMA_PATH = os.path.join(BASE_DIR, "data", "experience_schema.json")
    print(f"Current working directory: {os.getcwd()}")
    print(f"Looking for experience_schema.json at: {EXPERIENCE_SCHEMA_PATH}")  # Debugging

    try:
        with open(EXPERIENCE_SCHEMA_PATH, "r") as file:
            exp_schemas = json.load(file)
    except FileNotFoundError:
        raise FileNotFoundError(f"Experience schema file not found at {EXPERIENCE_SCHEMA_PATH}")

    current_date = datetime.datetime.now().date()

    # First get general experience
    messages = [
        SystemMessage(content=GENERAL_EXPERIENCE_SP),
        HumanMessage(
            content=(
                GENERAL_EXPERIENCE_HP.format(
                    text_input=pdf_text,
                    json_dump=json.dumps(
                        exp_schemas["general_experience"]["json_schema"], indent=2
                    ),
                )
                + f"\n\nCurrent Date: {current_date}"
            )
        ),
    ]

    response = llm.invoke(
        messages, functions=[exp_schemas["general_experience"]["json_schema"]]
    )
    structured_data = response.additional_kwargs["function_call"]["arguments"]
    general_exp_output_json = json.loads(structured_data)

    # Then use tailored experience prompt with job description
    messages = [
        SystemMessage(content=TAILORED_SEP_EXPERIENCE_SP),
        HumanMessage(
            content=(
                TAILORED_SEP_EXPERIENCE_HP.format(
                    text_input=json.dumps(general_exp_output_json, indent=2),
                    json_dump=json.dumps(
                        exp_schemas["sep_experience"]["json_schema"], indent=2
                    ),
                    job_description=job_description,
                    role=role,
                )
                + f"\n\nCurrent Date: {current_date}"
            )
        ),
    ]
    response = llm.invoke(
        messages, functions=[exp_schemas["sep_experience"]["json_schema"]]
    )
    structured_data = response.additional_kwargs["function_call"]["arguments"]
    json_structured_data = json.loads(structured_data)

    # Define the filename where you want to store the data
    filename = "structured_experience.json"

    modified_cgi_experience = []
    for entry in json_structured_data["cgi_experience"]:
        modified_entry = {f"cgi_{key}": value for key, value in entry.items()}
        modified_cgi_experience.append(modified_entry)
    json_structured_data["cgi_experience"] = modified_cgi_experience

    # Write the JSON data to the file
    with open(filename, "w") as file:
        json.dump(json_structured_data, file, indent=2)  # Use json.dump to format the data and write to the file

    return json_structured_data



def extract_text_from_docx_simple(file_path):
    """Simple text extraction from DOCX as fallback"""
    try:
        doc = Document(file_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            full_text.append(paragraph.text.strip())
        
        return '\n'.join(full_text)
    except Exception as e:
        return "Error extracting text from document"

def validate_and_fix_scores(evaluation_result):
    """
    Validate and fix common scoring issues in evaluation results
    """
    try:
        # Fix overall score if it's out of range (like 68/10 instead of 6.8/10)
        overall_score = evaluation_result.get('overall_score', 0)
        
        if overall_score > 10:
            # Likely a percentage, convert to 1-10 scale
            evaluation_result['overall_score'] = round(overall_score / 10, 1)
        elif overall_score < 0:
            evaluation_result['overall_score'] = 0
        
        # Fix individual section scores
        if 'ratings_summary' in evaluation_result:
            for section, score in evaluation_result['ratings_summary'].items():
                if score > 10:
                    evaluation_result['ratings_summary'][section] = round(score / 10, 1)
                elif score < 0:
                    evaluation_result['ratings_summary'][section] = 0
        
        # Recalculate overall score if needed
        if 'ratings_summary' in evaluation_result and 'section_weights' in evaluation_result:
            weights = evaluation_result['section_weights']
            ratings = evaluation_result['ratings_summary']
            
            calculated_score = sum(
                ratings.get(section, 0) * weight 
                for section, weight in weights.items()
            )
            
            # Only update if significantly different
            if abs(calculated_score - evaluation_result['overall_score']) > 0.5:
                evaluation_result['overall_score'] = round(calculated_score, 1)
        
        return evaluation_result
        
    except Exception as e:
        log(f"Error validating scores: {str(e)}")
        return evaluation_result

def run_quick_evaluation(llm, resume_text, file_path):
    """
    Fallback quick evaluation when main evaluation fails
    """
    try:
        from data.prompts import QUICK_EVALUATION_SP, QUICK_EVALUATION_HP
        
        messages = [
            SystemMessage(content=QUICK_EVALUATION_SP),
            HumanMessage(content=QUICK_EVALUATION_HP.format(resume_text=resume_text))
        ]
        
        response = llm.invoke(messages)
        
        # Try to parse the quick evaluation
        json_match = re.search(r'```json\n(.*?)\n```', response.content, re.DOTALL)
        if json_match:
            quick_result = json.loads(json_match.group(1))
        else:
            quick_result = json.loads(response.content.strip())
        
        
        standard_result = {
            "overall_score": quick_result.get('overall_score', 5.0),
            "critical_issues_identified": quick_result.get('critical_issues', []),
            "immediate_fixes": quick_result.get('immediate_fixes', []),
            "role_assessment": quick_result.get('role_assessment', {}),
            "experience_analysis": quick_result.get('experience_analysis', {}),
            "evaluation_type": "quick_fallback",
            "recommendations": quick_result.get('immediate_fixes', []),
            "evaluation_metadata": {
                "file_path": file_path,
                "evaluation_date": datetime.datetime.now().isoformat(),
                "fallback_used": True
            }
        }
        
        return standard_result
        
    except Exception as e:
        return {
            "overall_score": 0,
            "error": "All evaluation methods failed",
            "message": f"Evaluation error: {str(e)}",
            "evaluation_metadata": {
                "file_path": file_path,
                "evaluation_date": datetime.datetime.now().isoformat(),
                "complete_failure": True
            }
        }

def evaluate_resume(file_path, job_description=None, target_role=None, evaluation_criteria=None):
    """
    Enhanced resume evaluation with focus on specific issues like:
    - Recent vs older role detail levels
    - Project detail sufficiency  
    - Employment gaps
    - Role-specific alignment
    - Grammar and formatting
    
    Args:
        file_path (str): Path to the resume file (PDF or DOCX)
        job_description (str, optional): Target job description for tailored evaluation
        target_role (str, optional): Specific role type (e.g., 'Business Analyst', 'Developer', 'Director')
        evaluation_criteria (dict, optional): Custom evaluation criteria
        
    Returns:
        dict: Enhanced evaluation results with specific issue identification
    """
    try:
        load_dotenv()
        
        # Initialize LLM
        llm = AzureChatOpenAI(
            azure_endpoint=os.environ["AZURE_OPENAI_ENDPOINT"],
            api_key=os.environ["AZURE_OPENAI_API_KEY"],
            api_version="2024-12-01-preview",
            deployment_name="gpt-4o",
            model="gpt-4o",
        )
        
        # Extract resume data
        if file_path.endswith('.pdf'):
            loader = PyPDFLoader(file_path)
            pages = loader.load()
            resume_text = "\n".join([doc.page_content for doc in pages])
        elif file_path.endswith('.docx'):
            try:
                structured_data = extract_resume_data_with_llm(file_path, llm)
                resume_text = json.dumps(structured_data, indent=2)
            except Exception as docx_error:
                resume_text = extract_text_from_docx_simple(file_path)
        else:
            raise ValueError("Unsupported file format. Use PDF or DOCX.")
        
        if not resume_text or len(resume_text.strip()) < 100:
            raise ValueError(f"Insufficient text extracted from resume: {len(resume_text)} characters")
        
        # Prepare context for job description if provided
        job_description_context = ""
        if job_description:
            job_description_context = f"\nTarget Job Description:\n{job_description}\n\nEvaluate how well this resume aligns with the target role requirements."
        
        # Import the enhanced prompts
        from data.prompts import ENHANCED_EVALUATION_SP, ENHANCED_EVALUATION_HP, QUICK_EVALUATION_SP, QUICK_EVALUATION_HP
        
        
        evaluation_prompt = ENHANCED_EVALUATION_HP.format(
            resume_text=resume_text,
            job_description_context=job_description_context
        )
        system_prompt = ENHANCED_EVALUATION_SP
        
        # Generate evaluation
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=evaluation_prompt)
        ]
        
        response = llm.invoke(messages)
        
        # Parse JSON response
        try:
            json_match = re.search(r'```json\n(.*?)\n```', response.content, re.DOTALL)
            if json_match:
                evaluation_result = json.loads(json_match.group(1))
            else:
                evaluation_result = json.loads(response.content.strip())
            
            # Validate and fix scoring issues
            evaluation_result = validate_and_fix_scores(evaluation_result)
            
            
            evaluation_result["evaluation_metadata"] = {
                "file_path": file_path,
                "evaluation_date": datetime.datetime.now().isoformat(),
                "target_role": target_role,
                "has_job_description": bool(job_description),
                "evaluation_version": "enhanced_v2"
            }
                
            return evaluation_result
            
        except json.JSONDecodeError as e:
            log(f"JSON decode error, falling back to quick evaluation: {str(e)}")
            return run_quick_evaluation(llm, resume_text, file_path)
            
    except Exception as e:
        log(f"Evaluation error: {str(e)}")
        return {
            "error": str(e),
            "overall_score": 0,
            "message": "Evaluation failed",
            "evaluation_metadata": {
                "file_path": file_path,
                "evaluation_date": datetime.datetime.now().isoformat(),
                "error": True
            }
        }

def resume_stream(
    st,  # Not used - for compatibility
    progress_bar,  # ProgressTracker object
    base_progress,
    file_progress_weight,
    file_path,
    selected_format,
    custom_role_title="",
    job_description="",
    rfp_file_path=None,
    include_default_cgi=False,
):
    """Main resume processing function - adapted from Streamlit version"""
    load_dotenv()

    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    JSON_SCHEMA_PATH = os.path.join(BASE_DIR, "data", "json_schema.json")
    ALL_SCHEMAS_PATH = os.path.join(BASE_DIR, "data", "all_schemas.json")
    EXPERIENCE_SCHEMA_PATH = os.path.join(BASE_DIR, "data", "experience_schema.json")

    llm = AzureChatOpenAI(
        azure_endpoint=os.environ["AZURE_OPENAI_ENDPOINT"],
        api_key=os.environ["AZURE_OPENAI_API_KEY"],
        api_version="2024-12-01-preview",
        deployment_name="gpt-4o",
        model="gpt-4o",
    )

    # If RFP file is provided, generate a job description from it
    if rfp_file_path and os.path.exists(rfp_file_path):
        log(f"RFP file detected: {rfp_file_path}")
        # Determine role title to use for the RAG job description generation
        role_for_rag = (
            custom_role_title.strip() if custom_role_title else selected_format
        )

        # Generate job description from RFP
        rfp_job_description = generate_rag_job_description(
            llm, rfp_file_path, role_for_rag
        )

        # If successful and user didn't provide a job description, use the generated one
        if rfp_job_description and not (job_description and job_description.strip()):
            log("Using RFP-generated job description")
            job_description = rfp_job_description
        # If user provided a job description, combine it with the RFP-generated one
        elif rfp_job_description and job_description and job_description.strip():
            log("Combining user-provided job description with RFP-generated description")
            job_description = f"{job_description}\n\nAdditional requirements from RFP:\n{rfp_job_description}"

    # Update progress
    progress_bar.progress(base_progress + file_progress_weight * 0.1)

    # Handle both PDF and DOCX files
    if file_path.lower().endswith('.pdf'):
        loader = PyPDFLoader(file_path)
        pages = loader.load()
        pdf_text = "\n".join([doc.page_content for doc in pages])
    elif file_path.lower().endswith(('.docx', '.doc')):
        # Use the extract_text_from_docx function
        pdf_text = extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_path}")

    current_date = datetime.datetime.now().date()

    try:
        with open(JSON_SCHEMA_PATH, "r") as file:
            json_schema = json.load(file)
    except FileNotFoundError:
        raise FileNotFoundError(f"JSON schema file not found at {JSON_SCHEMA_PATH}")
   
    # Generate structured data
    structured_data = generate_llm_content(
        llm=llm,
        system_prompt=STRUCTURED_DATA_SP,
        human_prompt_template=STRUCTURED_DATA_HP,
        format_args={
            "pdf_text": pdf_text,
            "json_input": json.dumps(json_schema, indent=2),
        },
        functions=[json_schema],
        extract_function_call=True,
    )

    log("Completed Structured Data")
    
    # If custom_experiences is provided, add them to the structured_data experience array
    if custom_experiences and len(custom_experiences) > 0:
        log(f"Adding {len(custom_experiences)} custom experience entries to structured data")
        
        # Ensure experience array exists
        if "experience" not in structured_data:
            structured_data["experience"] = []
        
        # Add custom experiences to the experience array
        for exp in custom_experiences:
            experience_entry = {
                "company": exp.get("company", ""),
                "position_title": exp.get("position_title", ""),
                "start_date": exp.get("start_date", ""),
                "end_date": exp.get("end_date", ""),
                "responsibilities": [exp.get("description", "")],
                "technologies": []  # Could be extracted from description if needed
            }
            structured_data["experience"].append(experience_entry)
        
        log(f"Added {len(custom_experiences)} custom experiences to structured data")
    
    # Update progress - 20% complete for this file
    progress_bar.progress(base_progress + file_progress_weight * 0.2)

    custom_role_title_string = (
        f"\nAnd considering the target role type of {custom_role_title}, aligning with the general role type, while being more general:\n"
        if custom_role_title
        else ""
    )

    # Generate an appropriate role title if not provided by user
    if not custom_role_title or not custom_role_title.strip():
        role_title = generate_llm_content(
            llm=llm,
            system_prompt=ROLE_TITLE_GEN_SP,
            human_prompt_template=ROLE_TITLE_GEN_HP,
            format_args={
                "structured_data": json.dumps(structured_data, indent=2),
                "custom_role_title_string": custom_role_title_string,
            },
        )
        log(f"Generated role title: {role_title}")
    else:
        role_title = custom_role_title.strip()
        log(f"Using provided role title: {role_title}")

    # Generate tailored profile based on job description if provided
    if job_description and job_description.strip():
        profile = generate_llm_content(
            llm=llm,
            system_prompt=TAILORED_SUMMARY_SP,
            human_prompt_template=TAILORED_SUMMARY_HP,
            format_args={
                "structured_data": structured_data,
                "job_description": job_description,
                "role": role_title,
            },
            functions=[json_schema],
        )
        log(f"Generated tailored profile with job description for {role_title}")
    else:
        profile = generate_llm_content(
            llm=llm,
            system_prompt=SUMMARY_SP,
            human_prompt_template=SUMMARY_HP,
            format_args={"structured_data": structured_data},
            functions=[json_schema],
        )
        log("Generated standard profile")

    # Generate years of experience
    years_exp = generate_llm_content(
        llm=llm,
        system_prompt=PROFILE_SP,
        human_prompt_template=PROFILE_HP,
        format_args={"profile": profile},
    )

    # Define function to call LLM with appropriate schema
    def call_llm(all_schemas, section, text_input=pdf_text, job_desc=""):
        sp_var = f"{section.upper()}_SP"
        hp_var = f"{section.upper()}_HP"

        # Only use tailored prompts if job description is provided and not empty
        if (
            job_desc
            and job_desc.strip()
            and f"TAILORED_{section.upper()}_SP" in globals()
        ):
            sp_var = f"TAILORED_{section.upper()}_SP"
            hp_var = f"TAILORED_{section.upper()}_HP"

        sp = eval(sp_var)
        hp = eval(hp_var)

        # Add job description to format args if available and needed
        format_args = {
            "text_input": text_input,
            "json_dump": json.dumps(all_schemas[section]["json_schema"], indent=2),
        }

        if (
            job_desc
            and job_desc.strip()
            and f"TAILORED_{section.upper()}_HP" in globals()
        ):
            format_args["job_description"] = job_desc
            format_args["role"] = role_title

        messages = [
            SystemMessage(content=sp),
            HumanMessage(content=hp.format(**format_args)),
        ]

        response = llm.invoke(messages, functions=[all_schemas[section]["json_schema"]])
        structured_data = response.additional_kwargs["function_call"]["arguments"]
        json_structured_data = json.loads(structured_data)

        return json_structured_data

    try:
        with open(ALL_SCHEMAS_PATH, "r") as file:
            all_schemas = json.load(file)
    except FileNotFoundError:
        raise FileNotFoundError(f"All schemas file not found at {ALL_SCHEMAS_PATH}")

    res_dict = dict()

    log("Loading...")
    # Process experience with job description if available and not empty
    if job_description and job_description.strip():
        res_dict["experience"] = call_tailored_experience_chain(
            pdf_text, job_description, role_title, llm
        )
        log(f"\t>> Completed key: experience (tailored for {role_title})")
    else:
        res_dict["experience"] = experience_chain(pdf_text, llm)
        log(f"\t>> Completed key: experience")

    # If include_default_cgi is True, generate a default CGI experience entry
    if include_default_cgi:
        log("Generating default CGI experience entry")
        default_cgi_exp = generate_cgi_experience(
            llm, selected_format, custom_role_title
        )

        # Add the default entry to the beginning of the CGI experience array
        if "cgi_experience" in res_dict["experience"]:
            # Check if there's already a non-empty cgi_experience
            if (
                res_dict["experience"]["cgi_experience"]
                and len(res_dict["experience"]["cgi_experience"]) > 0
            ):
                # Check if the first entry is a placeholder (client descriptor not provided)
                if res_dict["experience"]["cgi_experience"][0].get(
                    "cgi_client_or_sector"
                ) == "client descriptor not provided" or not res_dict["experience"][
                    "cgi_experience"
                ][
                    0
                ].get(
                    "cgi_responsibilities"
                ):
                    # Replace the placeholder with our generated experience
                    res_dict["experience"]["cgi_experience"][0] = default_cgi_exp
                    log("Replaced placeholder CGI experience with generated experience")
                else:
                    # Insert at the beginning if there's already valid content
                    res_dict["experience"]["cgi_experience"].insert(0, default_cgi_exp)
                    log("Added generated CGI experience at the beginning of existing experiences")
            else:
                # If the array is empty or None, initialize it with our experience
                res_dict["experience"]["cgi_experience"] = [default_cgi_exp]
                log("Created new CGI experience entry")

        log("Added default CGI experience entry")

    # Process other sections
    for key in all_schemas.keys():
        res_dict[key] = call_llm(all_schemas, key, pdf_text, job_description)
        log(f"\t>> Completed key: {key}")
    progress_bar.progress(base_progress + file_progress_weight * 0.4)

    # For testing
    data = {
        "structured_data": structured_data,
        "years_exp": years_exp,
        "profile": profile,
        "res_dict": res_dict,
        "job_description": job_description if job_description else "None provided",
    }
    with open("resume_data.json", "w") as json_file:
        json.dump(data, json_file, indent=4)

    print("Data saved to resume_data.json")

    # Pass job description, role title, and format type to generate_resume
    generate_resume(
        structured_data,
        years_exp,
        profile,
        res_dict,
        job_description,
        role_title,
        selected_format,
    )

def run_resume_evaluation(file_path, file_id, resume_name="Resume", selected_format=None, custom_role_title=None, job_description=None):
    """
    Run the enhanced resume evaluation with comprehensive error handling
    """
    try:
        # Verify file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Resume file not found: {file_path}")
        
        
        target_role = None
        
        # Run enhanced evaluation
        evaluation_results = evaluate_resume(
            file_path=file_path,
            job_description=job_description if job_description and job_description.strip() else None,
            target_role=target_role
        )
        
        # Always save JSON backup
        eval_json_path = f"evaluation_{file_id}.json"
        try:
            with open(eval_json_path, 'w', encoding='utf-8') as f:
                json.dump(evaluation_results, f, indent=2, ensure_ascii=False)
        except Exception as json_error:
            pass
        
        # Try to generate PDF
        eval_pdf_path = f"evaluation_{file_id}.pdf"
        try:
            from src.pdf_evaluation_generator import generate_evaluation_pdf, REPORTLAB_AVAILABLE
            
            if REPORTLAB_AVAILABLE:
                generate_evaluation_pdf(evaluation_results, eval_pdf_path, resume_name)
                
                if os.path.exists(eval_pdf_path):
                    return eval_pdf_path, evaluation_results
                else:
                    return eval_json_path, evaluation_results
            else:
                return eval_json_path, evaluation_results
                
        except ImportError as import_error:
            return eval_json_path, evaluation_results
        except Exception as pdf_error:
            return eval_json_path, evaluation_results
            
    except Exception as e:
        return None, None