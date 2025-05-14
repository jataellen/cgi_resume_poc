from spire.doc import *
from spire.doc.common import *
from docx import Document as DocxDocument
import io


# Function to convert a single DOCX file to PDF using Spire.Doc
def convert_docx_to_pdf(uploaded_file, temp_file_path, pdf_file_path):
    # Save the uploaded DOCX file temporarily
    with open(temp_file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    print(f"File saved as: {temp_file_path}")

    document = Document()
    document.LoadFromFile(temp_file_path)

    document.SaveToFile(pdf_file_path, FileFormat.PDF)
    document.Close()

    return pdf_file_path


def change_all_fonts(uploaded_file):
    # Create a BytesIO object from the uploaded file's buffer
    file_buffer = io.BytesIO(uploaded_file.getbuffer())

    # Load the document with python-docx
    doc = DocxDocument(file_buffer)

    # Change font for all runs in all paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"

    # Change font in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Arial"

    # Change font in headers and footers
    for section in doc.sections:
        for header in section.header.paragraphs:
            for run in header.runs:
                run.font.name = "Arial"

        for footer in section.footer.paragraphs:
            for run in footer.runs:
                run.font.name = "Arial"

    # Save the modified document to a new BytesIO object
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)  # Reset the buffer position to the beginning

    # Create a new FileStorage-like object with the modified document
    class ModifiedFile:
        def __init__(self, file_buffer, orig_file):
            self.buffer = file_buffer
            self.type = orig_file.type

        def getbuffer(self):
            return self.buffer.getvalue()

        def read(self):
            return self.buffer.read()

    return ModifiedFile(output_buffer, uploaded_file)


def convert_to_pdf(uploaded_file, file_id):
    temp_file_path = f"../temp_{file_id}.docx"
    pdf_file_path = f"../temp_{file_id}.pdf"

    if (
        uploaded_file.type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        # Change all fonts to Arial
        modified_file = change_all_fonts(uploaded_file)

        # Convert to PDF
        pdf_file_path = convert_docx_to_pdf(
            modified_file, temp_file_path, pdf_file_path
        )
    else:
        with open(pdf_file_path, "wb") as f:
            f.write(uploaded_file.read())

    return pdf_file_path
