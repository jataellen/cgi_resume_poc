from src.logs_manager import log

from docx import Document

def replace_text_in_docx(doc: Document, replacements: list):
    
    for key, value in replacements:
         # Determine what to display
        if not value or (isinstance(value, list) and not any(item.strip() for item in value)):
            value = f"[Missing field: {key.strip('{}')}]"

        def process_paragraphs(paragraphs):
            for paragraph in paragraphs:
                if key in paragraph.text:
                    if isinstance(value, list):
                        # Remove the placeholder 
                        # idx : Index of the placeholder
                        parent = paragraph._element.getparent()
                        idx = parent.index(paragraph._element)
                        parent.remove(paragraph._element)

                        # Insert bullet points
                        for item in value:
                            if item and item.strip():  
                                bullet_para = doc.add_paragraph(item.strip(), style="ListBullet")
                                parent.insert(idx, bullet_para._element)
                                idx += 1
                        if not any(item.strip() for item in value):
                            # Insert Missing field if list is empty
                            missing_para = doc.add_paragraph(f"[Missing field: {key.strip('{}')}]", style="Normal")
                            parent.insert(idx, missing_para._element)
                    else:
                        # replacing the content if it is not a list
                        paragraph.text = paragraph.text.replace(key, str(value))
                    return True
            return False

        # Try replacing in normal document paragraphs
        if process_paragraphs(doc.paragraphs):
            continue

        # Try replacing inside table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if process_paragraphs(cell.paragraphs):
                        break