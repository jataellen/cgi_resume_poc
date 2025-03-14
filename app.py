import streamlit as st
import pandas as pd
import numpy as np
import os
from openai import AzureOpenAI
from langchain_openai import AzureChatOpenAI
from langchain.schema import SystemMessage, HumanMessage
import json
from langchain_community.document_loaders import PyPDFLoader
import datetime
from docx import Document

os.environ["AZURE_OPENAI_API_KEY"] = "EOkfcf05uMhPPi5vtu0OmXUMrpdNc4Ji65zbVs1iZZGbbdGvunPhJQQJ99BBACYeBjFXJ3w3AAABACOGejoY"
os.environ["AZURE_OPENAI_ENDPOINT"] = "https://cgi-resume-openai.openai.azure.com/openai/deployments/gpt-4o/chat/completions?api-version=2024-10-21"

def replicate_section(doc, start_tag, end_tag, replacements, times_to_repeat):
    """
    Duplicates the section between start_tag and end_tag (inclusive), replacing the tags with specified values.
    """
   
    para_group = []
    inside_section = False
    end_para = None
    
    for paragraph in doc.paragraphs:
        if start_tag in paragraph.text:
            inside_section = True
            paragraph.text = paragraph.text.replace(f"{start_tag}", "")
        if inside_section:
            para_group.append(paragraph)
        if end_tag in paragraph.text and inside_section:
            end_para = paragraph
            inside_section = False
            paragraph.text = paragraph.text.replace(f"{end_tag}", "")
            break 
    
    end_para = end_para._p
    for i in range(times_to_repeat):
        new_para_lst = []
        for paragraph in para_group:
            paragraph.text = paragraph.text.replace(f"{start_tag}", "")
            paragraph.text = paragraph.text.replace(f"{end_tag}", "")
            
            new_paragraph = doc.add_paragraph()
            new_paragraph.alignment = paragraph.alignment
            new_paragraph.style = paragraph.style
            

            new_paragraph.text = paragraph.text    
            new_para_lst.append(new_paragraph)
            
        for para in new_para_lst:
            end_para.addnext(para._p)
            end_para = para._p

# def replicate_row(doc, key, res_dict_skills):
    




def replace_text_in_docx(doc, replacements):
    """
    Replaces text in a DOCX file, and replicates sections based on the specified start and end tags.
    """
    filtered_replacements = {k: v for k, v in replacements.items() if v is not None}
    for key, value in filtered_replacements:
        log(f"Working on key: {key}")
        # st.write(f"Working on key: {key}\nValue: {value}")
        for paragraph in doc.paragraphs:
        
            if key in paragraph.text:
                if isinstance(value, list):
                    paragraph.style = 'ListBullet'
                    x_par = paragraph._p
                    if len(value) > 1:
                        paragraph.text = paragraph.text.replace(key, value[0])
                        value = value[1:]
                        for bp in value[::-1]:
                            para = doc.add_paragraph(bp, style='ListBullet')
                            x_par.addnext(para._p)
                else:
                    paragraph.text = paragraph.text.replace(key, value)
                break
                
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if key in paragraph.text:
                            if isinstance(value, list):
                                paragraph.style = 'ListBullet'
                                x_par = paragraph._p
                                paragraph.text = paragraph.text.replace(key, value[0])
                                value = value[1:]
                                for bp in value:
                                    para = doc.add_paragraph(bp, style='ListBullet')
                                    x_par.addnext(para._p)
                            else:
                                paragraph.text = paragraph.text.replace(key, value)
                            break
                        

def handle_skills_summary(doc, replacements, res_dict_skills):
    for key, value in replacements:
        for table in doc.tables:
            for row in table.rows:
                if any(key in para.text for cell in row.cells for para in cell.paragraphs):
                    cleaned_key = key.replace("{", "").replace("}", "")
                    times_to_repeat = len(res_dict_skills[cleaned_key]) - 1
                    for _ in range(times_to_repeat):
                        new_row = table.add_row()
                        for i, new_cell in enumerate(new_row.cells):
                            old_para = row.cells[i].paragraphs[0]
                            new_cell.text = old_para.text
                            new_cell.paragraphs[0].style = old_para.style
                        row._tr.addnext(new_row._tr)


def replace_text_in_table(doc, replacements, res_dict):
    """
    Replaces text in a DOCX file, and replicates sections based on the specified start and end tags.
    """
    for key, value in replacements:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replacement_made = False
                    for paragraph in cell.paragraphs:
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)
                            replacement_made = True
                            break
                    if replacement_made:
                        break  
                if replacement_made:
                    break 
            if replacement_made:
                break  

def generate_resume(structured_data, years_exp, profile, res_dict):
    # Input/Output
    input_filename = "resume_sample.docx"
    output_filename = "updated_resume.docx"
    doc = Document(input_filename)
    
    full_name = structured_data['contact']['name']
    cgi_title = "Consultant"
    sector = "Health Services"
    certs = None
    if 'certifications' in structured_data:
        certs = [f"{i['name']}, {i['issuing_organization']}" for i in structured_data['certifications']]
    replacements = [
        ("{full_name}", full_name),
        ("{cgi_title}", cgi_title),
        ("{years_exp}", years_exp),
        ("{professional_profile}", profile),
        ("{industry}", res_dict['other_sections']['industry_experience']),
        ("{tech_specs}", res_dict['other_sections']['technical_specializations']),
        ("{expertise}", res_dict['other_sections']['areas_of_expertise']),
        ("{languages}", res_dict['other_sections']['languages']),
        ("{environment}", res_dict['other_sections']['environments']),
        ("{tools}", res_dict['other_sections']['tools_and_software']),
        ("{certs}", certs),
    ]
    
    # if 'certifications' in structured_data:

    
    # CGI Experience
    cgi_exp =  res_dict['experience']['cgi_experience']
    for exp in cgi_exp:
        exp = {k: v for k, v in exp.items() if k in ['sector', 'job_title', 'start_date', 'end_date', 'responsibilities']}
        for key, value in exp.items():
            replacements.append( ("{" + key + "}", value) )
    times_to_repeat = len(cgi_exp) -1
    
    replicate_section(doc, "{begin_cgi_exp}", "{end_cgi_exp}", replacements, times_to_repeat)

    # Other Experience
    o_exp =  res_dict['experience']['other_experience']
    for exp in o_exp:
        exp = {k: v for k, v in exp.items() if k in ['company', 'job_title', 'start_date', 'end_date', 'responsibilities']}
        for key, value in exp.items():
            replacements.append( ("{" + key + "}", value) )
    times_to_repeat = len(o_exp) -1
    replicate_section(doc, "{begin_other_exp}", "{end_other_exp}", replacements, times_to_repeat)

    # Skills summary
    table_reps = []
    for key, value in res_dict['skills_summary'].items():
        table_reps.append( ("{" + key + "}", value) )
    ed_list = [f"{el['degree']}, {el['field_of_study']} - {el['institution']}" for el in structured_data['education']]
    replacements.append( ("{education_entry}", ed_list) )


    replace_text_in_docx(doc, replacements)
    handle_skills_summary(doc, table_reps, res_dict['skills_summary'])

    replacements = []
    for key, value in res_dict['skills_summary'].items():
        replacements.extend( [ ("{" + key + "}", v['skill']) for v in value] )
        replacements.extend( [ ("{num_years}", str(v['years_of_experience'])) for v in value] )
        replacements.extend( [ ("{skill_level}", str(v['skill_level'])) for v in value] )

    replace_text_in_table(doc, replacements,  res_dict['skills_summary'])

    
    doc.save(output_filename)
    print(f"Updated document saved as: {output_filename}")



def log(message):
    log_messages.append(message)
    log_box.text_area("Logs", value="\n".join(log_messages), height=200)

log_messages = []
# Create a placeholder for the log box
log_box = None

def resume_stream(st, file_path):
    llm = AzureChatOpenAI(
        azure_endpoint=os.environ["AZURE_OPENAI_ENDPOINT"],
        api_key=os.environ["AZURE_OPENAI_API_KEY"],
        api_version="2024-12-01-preview",
        deployment_name="gpt-4o",
        model="gpt-4o",  # Ensure function calling support
    )

    # file_path = "Bila Gaite Resume - CGI Consulting  copy 1.pdf"

    loader = PyPDFLoader(file_path)
    pages = []
    # async for page in loader.alazy_load():
    #     pages.append(page)
    pages = loader.load()

    pdf_text = "\n".join([doc.page_content for doc in  pages])

    current_date = datetime.datetime.now().date()

    with open('json_schema.json', 'r') as file:
        json_schema = json.load(file)

    messages = [
        SystemMessage(
            content="You are an AI that extracts structured information from plain text resumes and returns JSON output."
        ),
        HumanMessage(
            content=(
                f"Extract and structure the following text into JSON format:\n\n{pdf_text}\n\n"
                f"Ensure the response matches this schema:\n{json.dumps(json_schema, indent=2)}"
            )
        ),
    ]

    response = llm.invoke(messages, functions=[json_schema])
    structured_data = response.additional_kwargs["function_call"]["arguments"]
    structured_data = json.loads(structured_data)
    
    log("Completed Structured Data")
    # return structured_data

    ### Professional Summary
    messages = [
        SystemMessage(
            content="You are an AI that takes structured resumes in JSON format and writes a compelling, professional summary of the applicant."
        ),
        HumanMessage(
            content=(
                "Using the following structured resume data in JSON format:\n\n"
                f"{structured_data}\n\n"
                "Write a well-crafted, three-paragraph professional profile of the applicant in the third person. "
                "Keep a good balance of detailed and concise. Do not use AI-isms"
                "Incorporate their professional summary, work experience, education, skills, certifications, and any notable achievements. "
                "Highlight their expertise, impact, and technical skills, ensuring the profile flows naturally and is engaging."
            )
        ),
    ]


    response = llm.invoke(messages, functions=[json_schema])

    profile = response.content.strip()

    # Years of Experience
    messages = [
        SystemMessage(
            content="You are an AI that takes in a professional summary and determines the applicants years of experience."
        ),
        HumanMessage(
            content=(
                "Using the following professional summary:\n\n"
                f"{profile}\n\n"
                "Write a very concise header desribing their experience in the following format:\n <X> years experience in <X_category>\nex: 5 years of experience in Software Development"
            )
        ),
    ]


    # Call the LLM with function calling enabled
    response = llm.invoke(messages, functions=[json_schema])

    years_exp = response.content.strip()

    def call_llm(overall, section, text_input=pdf_text):
        messages = [
            SystemMessage(
                content=overall[section]['system_prompt']
            ),
            HumanMessage(
                content=(overall[section]['human_prompt'].format(text_input=text_input, json_dump=json.dumps(overall[section]['json_schema'], indent=2)))
            ), 
        ]
        response = llm.invoke(messages, functions=[overall[section]['json_schema']])
        structured_data = response.additional_kwargs["function_call"]["arguments"]
        json_structured_data = json.loads(structured_data)
        
        return json_structured_data
    
    with open("overall.json", 'r') as file:
        overall = json.load(file)
    
    res_dict = dict()

    log("Loading...")
    for key in overall.keys():
        res_dict[key] = call_llm(overall, key, pdf_text)
        log(f"\t>> Completed key: {key}")

    generate_resume(structured_data, years_exp, profile, res_dict)




st.title('ResumeGenie')

with st.form("my-form", clear_on_submit=True):
    uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")
    submitted = st.form_submit_button("Submit")



log_box = st.empty()

if submitted and uploaded_file is not None:
    # Save the uploaded file temporarily
    original_file_name = uploaded_file.name
    with open("temp.pdf", "wb") as f:
        f.write(uploaded_file.read())
    log( f"Uploaded file: {original_file_name}")

    # Process the PDF (assuming resume_stream creates updated_resume.docx)
    resume_stream( st, "temp.pdf")
    log("Processed the PDF")

    # Rename the updated resume, and delete if it already exists
    new_file_name = os.path.splitext(original_file_name)[0] + "_updated.docx"
    if os.path.exists(new_file_name):
        os.remove(new_file_name)
        log(f"Deleted existing file: {new_file_name}")
    os.rename("updated_resume.docx", new_file_name)
    log( f"Renamed updated resume to: {new_file_name}")

    # Provide a download link for the updated resume
    with open(new_file_name, "rb") as f:
        st.download_button(label="Download Updated Resume", data=f, file_name=new_file_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    log("Provided download link for updated resume")


    

else:
    st.write("No PDF uploaded yet.")

